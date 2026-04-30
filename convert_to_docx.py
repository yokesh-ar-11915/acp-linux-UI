"""Convert Linux-ACP-User-Scenarios.md to a rich .docx document."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MD_PATH = Path(r"d:\linux-acp\acp-dashboard\Linux-ACP-User-Scenarios.md")
OUT_PATH = Path(r"d:\linux-acp\acp-dashboard\Linux-ACP-User-Scenarios.docx")

# ── Styles & Colours ──────────────────────────────────────────────
CLR_PRIMARY   = RGBColor(0x1A, 0x73, 0xE8)  # ManageEngine blue
CLR_DARK      = RGBColor(0x22, 0x22, 0x22)
CLR_GREY      = RGBColor(0x55, 0x55, 0x55)
CLR_LIGHT_BG  = RGBColor(0xF5, 0xF6, 0xFA)
CLR_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
CLR_TABLE_HDR = RGBColor(0x1A, 0x73, 0xE8)

FONT_NAME = "Segoe UI"
FONT_MONO = "Consolas"


def setup_styles(doc):
    """Configure document styles."""
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)
    style.font.color.rgb = CLR_DARK
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.15

    for level, (size, bold, color, before, after) in {
        "Heading 1": (18, True, CLR_PRIMARY, 18, 8),
        "Heading 2": (14, True, CLR_DARK, 14, 6),
        "Heading 3": (12, True, CLR_DARK, 10, 4),
    }.items():
        s = doc.styles[level]
        s.font.name = FONT_NAME
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True


def set_cell_shading(cell, color_hex):
    """Set cell background colour."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows):
    """Add a table with styled header row."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = CLR_WHITE
        run.font.name = FONT_NAME
        set_cell_shading(cell, "1A73E8")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            run.font.name = FONT_NAME
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F5F6FA")

    # Column widths — distribute evenly
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)

    return table


def add_code_block(doc, text):
    """Add a monospaced code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = FONT_MONO
    run.font.size = Pt(8.5)
    run.font.color.rgb = CLR_DARK
    # Background shading on the paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
    pPr.append(shading)


def add_rich_paragraph(doc, text, style=None, bold=False, italic=False,
                       font_size=None, color=None, alignment=None, space_before=None, space_after=None):
    """Add a paragraph with inline bold/**text**/ handling."""
    p = doc.add_paragraph(style=style)
    if alignment:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)

    # Split on **bold** and `code` markers
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = FONT_MONO
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            run = p.add_run(part)
            if bold:
                run.bold = True
            if italic:
                run.italic = True

        run.font.name = run.font.name or FONT_NAME
        if font_size:
            run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = color

    return p


def add_horizontal_rule(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def parse_table_block(lines):
    """Parse markdown table lines into (headers, rows)."""
    headers = []
    rows = []
    for i, line in enumerate(lines):
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if i == 0:
            headers = cells
        elif i == 1:
            continue  # separator row
        else:
            # Pad to match header count
            while len(cells) < len(headers):
                cells.append("")
            rows.append(cells[:len(headers)])
    return headers, rows


def convert():
    md_text = MD_PATH.read_text(encoding="utf-8")
    lines = md_text.splitlines()

    doc = Document()
    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    setup_styles(doc)

    # ── Title block ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("HIGH LEVEL USE CASE DOCUMENT")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = CLR_PRIMARY
    run.font.name = FONT_NAME

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Application Control — ManageEngine Endpoint Central")
    run.font.size = Pt(13)
    run.font.color.rgb = CLR_GREY
    run.font.name = FONT_NAME

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Date: April 7, 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = CLR_GREY
    run.font.name = FONT_NAME

    add_horizontal_rule(doc)

    # ── Parse the rest of the document ──
    i = 5  # Skip the first 4 lines (title block + ---)
    n = len(lines)
    in_code_block = False
    code_lines = []

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Skip standalone --- separators
        if stripped == "---":
            add_horizontal_rule(doc)
            i += 1
            continue

        # Code blocks
        if stripped.startswith("```"):
            if in_code_block:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue

        # Tables — collect all consecutive | lines
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                headers, rows = parse_table_block(table_lines)
                if headers and rows:
                    add_styled_table(doc, headers, rows)
                    doc.add_paragraph()  # spacing after table
            continue

        # Numbered list items
        m = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if m:
            add_rich_paragraph(doc, stripped, font_size=10)
            i += 1
            continue

        # Bullet points (- or *)
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            # Indented sub-items
            indent_level = len(line) - len(line.lstrip())
            p = add_rich_paragraph(doc, f"• {text}", font_size=10)
            if indent_level > 0:
                p.paragraph_format.left_indent = Cm(0.5 * (indent_level // 2 + 1))
            else:
                p.paragraph_format.left_indent = Cm(0.5)
            i += 1
            continue

        # Bold-only lines (like **Milestone:** ...)
        if stripped.startswith("**") and "**" in stripped[2:]:
            add_rich_paragraph(doc, stripped, font_size=10)
            i += 1
            continue

        # Empty lines
        if not stripped:
            i += 1
            continue

        # Regular paragraphs
        add_rich_paragraph(doc, stripped, font_size=10)
        i += 1

    # Footer
    add_horizontal_rule(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Document Version: 2.0 | Created: April 7, 2026 | Scope: Linux Application Control Policy Deployment — M1")
    run.font.size = Pt(8)
    run.font.color.rgb = CLR_GREY
    run.font.italic = True
    run.font.name = FONT_NAME

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Format: High-Level Use Case Document per ManageEngine Endpoint Central standards")
    run.font.size = Pt(8)
    run.font.color.rgb = CLR_GREY
    run.font.italic = True
    run.font.name = FONT_NAME

    doc.save(str(OUT_PATH))
    print(f"✓ Saved: {OUT_PATH}")


if __name__ == "__main__":
    convert()
